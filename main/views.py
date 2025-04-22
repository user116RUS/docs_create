from django.core.paginator import Paginator
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.http import HttpResponse, FileResponse, JsonResponse
from django.template.loader import get_template
import os
import io
from docx import Document as DocxDocument
from docx.shared import Pt
from docxtpl import DocxTemplate
from django.forms import modelformset_factory
from django.utils.decorators import method_decorator
import tempfile
from openai import OpenAI
import re

from main.models import Document, Organisation, Service, ViewerCategory

from .forms import DocsForm, OrganisationForm, ServiceForm, ViewerCategoryForm, ViewerCategoryFormset

from django.views.generic.edit import CreateView, UpdateView
from django.views.generic.list import ListView
import json


def index(request):
    print('> tested')
    return render(request, template_name='main/index.html')


@login_required
def organisations_list(request):
    """Функция для отображения списка организаций-контрагентов"""
    # Получаем все организации
    organisations = Organisation.objects.all().order_by('name')
    
    # Добавляем пагинацию
    paginator = Paginator(organisations, 10)  # 10 организаций на страницу
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'main/organisations_list.html', {
        'page_obj': page_obj,
        'count': organisations.count(),
    })


@method_decorator(login_required, name='dispatch')
class OrganisationUpdateView(UpdateView):
    """Класс для редактирования организации"""
    model = Organisation
    form_class = OrganisationForm
    template_name = 'main/edit_organisation.html'
    
    def get_success_url(self):
        return reverse('main:organisations_list')


@login_required
def delete_organisation(request, pk):
    """Функция для удаления организации"""
    organisation = get_object_or_404(Organisation, pk=pk)
    
    if request.method == 'POST':
        # Проверяем, используется ли организация в документах
        documents_as_customer = Document.objects.filter(customer=organisation).count()
        documents_as_doer = Document.objects.filter(doer=organisation).count()
        
        if documents_as_customer > 0 or documents_as_doer > 0:
            # Если организация используется, возвращаем ошибку
            return JsonResponse({
                'success': False,
                'error': f'Невозможно удалить организацию, так как она используется в документах: {documents_as_customer + documents_as_doer} шт.'
            })
        else:
            # Если не используется, удаляем
            organisation.delete()
            return JsonResponse({'success': True})
    
    # Если запрос не POST, возвращаем ошибку
    return JsonResponse({'success': False, 'error': 'Метод не поддерживается'})


@login_required
def documents_list(request):
    """Функция для отображения списка всех документов"""
    # Получаем все документы
    documents = Document.objects.all().order_by('-id')
    
    # Поиск по параметрам
    search_query = request.GET.get('q', '')
    if search_query:
        documents = documents.filter(
            act_and_account_number__icontains=search_query) | documents.filter(
            number_basis_of_the_contract__icontains=search_query) | documents.filter(
            customer__name__icontains=search_query) | documents.filter(
            doer__name__icontains=search_query)
    
    # Фильтр по дате
    date_filter = request.GET.get('date', '')
    if date_filter:
        documents = documents.filter(date=date_filter)
        
    # Добавляем пагинацию
    paginator = Paginator(documents, 10)  # 10 документов на страницу
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'main/documents_list.html', {
        'page_obj': page_obj,
        'count': documents.count(),
        'search_query': search_query,
        'date_filter': date_filter
    })


@method_decorator(login_required, name='dispatch')
class DocumentUpdateView(UpdateView):
    """Класс для редактирования документа"""
    model = Document
    form_class = DocsForm
    template_name = 'main/edit_document.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        services = Service.objects.all()
        context['services'] = services
        return context
    
    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        form = self.get_form()
        
        # Проверяем, что форма валидна
        if form.is_valid():
            try:
                # Создаем и сохраняем объект
                self.object = form.save(commit=False)
                
                # Автозаполнение поля address_and_time
                if 'customer' in form.cleaned_data and form.cleaned_data['customer'] and 'time' in form.cleaned_data:
                    customer = form.cleaned_data['customer']
                    time = form.cleaned_data['time']
                    # Формируем строку "адрес. Время проведения: с [время]"
                    if customer.address and time:
                        self.object.address_and_time = f"{customer.address}. Время проведения: с {time.split('-')[0].strip()}"
                
                # Сохраняем объект
                self.object.save()
                # Сохраняем связи many-to-many
                form.save_m2m()
                
                # Явно указываем перенаправление на страницу загрузки документов
                success_url = reverse('main:document_detail', kwargs={'document_id': self.object.id})
                return redirect(success_url)
            except Exception as e:
                form.add_error(None, f"Ошибка при сохранении документа: {e}")
                return self.form_invalid(form)
        else:
            # Проверяем, есть ли услуги в запросе
            if 'services' not in request.POST and not request.POST.getlist('services'):
                form.add_error(None, "Пожалуйста, выберите хотя бы одну услугу.")
                
            # Проверяем, заполнено ли поле количества зрителей
            if not request.POST.get('viewers', '').strip():
                form.add_error('viewers', "Пожалуйста, укажите количество зрителей.")
                
            # Проверяем, заполнено ли поле суммы прописью
            if not request.POST.get('price_in_words', '').strip():
                form.add_error('price_in_words', "Пожалуйста, укажите сумму прописью.")
            
            return self.form_invalid(form)
    
    def get_success_url(self):
        return reverse('main:document_detail', kwargs={'document_id': self.object.id})


@method_decorator(login_required, name='dispatch')
class CreateDocs(CreateView):
    form_class = DocsForm
    template_name = 'main/create_docs.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        services = Service.objects.all()
        
        # Добавляем отладочную информацию
        for service in services:
            print(f"Service: {service.name}")
            print(f"Total viewers: {service.total_viewers}")
            print(f"Total price: {service.total_price}")
            print("Categories:")
            for category in service.viewer_categories.all():
                print(f"  - {category.viewers} viewers, {category.price} rub")
            print("---")
        
        context['services'] = services
        return context
    
    def post(self, request, *args, **kwargs):
        self.object = None
        form = self.get_form()
        
        # Проверяем, что форма валидна
        if form.is_valid():
            try:
                # Создаем и сохраняем объект
                self.object = form.save(commit=False)
                
                # Автозаполнение поля address_and_time
                if 'customer' in form.cleaned_data and form.cleaned_data['customer'] and 'time' in form.cleaned_data:
                    customer = form.cleaned_data['customer']
                    time = form.cleaned_data['time']
                    # Формируем строку "адрес. Время проведения: с [время]"
                    if customer.address and time:
                        self.object.address_and_time = f"{customer.address}. Время проведения: с {time.split('-')[0].strip()}"
                
                # Сохраняем объект
                self.object.save()
                # Сохраняем связи many-to-many
                form.save_m2m()
                
                print(f"Document created: {self.object.id}")
                
                # Явно указываем перенаправление на страницу загрузки документов
                success_url = reverse('main:document_detail', kwargs={'document_id': self.object.id})
                return redirect(success_url)
            except Exception as e:
                print(f"Error saving document: {e}")
                form.add_error(None, f"Ошибка при сохранении документа: {e}")
                return self.form_invalid(form)
        else:
            # Проверяем, есть ли услуги в запросе
            if 'services' not in request.POST and not request.POST.getlist('services'):
                form.add_error(None, "Пожалуйста, выберите хотя бы одну услугу.")
                
            # Проверяем, заполнено ли поле количества зрителей
            if not request.POST.get('viewers', '').strip():
                form.add_error('viewers', "Пожалуйста, укажите количество зрителей.")
                
            # Проверяем, заполнено ли поле суммы прописью
            if not request.POST.get('price_in_words', '').strip():
                form.add_error('price_in_words', "Пожалуйста, укажите сумму прописью.")
            
            print(f"Form is invalid: {form.errors}")
            return self.form_invalid(form)
    
    def get_success_url(self):
        return reverse('main:document_detail', kwargs={'document_id': self.object.id})


@method_decorator(login_required, name='dispatch')
class OrganisationView(CreateView):
    form_class = OrganisationForm
    template_name = 'main/create_organisation.html'
    success_url = '/'


@login_required
def create_service(request):
    if request.method == 'POST':
        form = ServiceForm(request.POST)
        if form.is_valid():
            service = form.save()
            
            # Обработка формсета с категориями зрителей
            formset = ViewerCategoryFormset(request.POST, instance=service)
            if formset.is_valid():
                formset.save()
                return redirect('main:create_docs')
    else:
        form = ServiceForm()
        formset = ViewerCategoryFormset()
    
    return render(request, 'main/create_service.html', {
        'form': form,
        'formset': formset
    })


@login_required
def edit_service(request, service_id):
    service = get_object_or_404(Service, id=service_id)
    
    if request.method == 'POST':
        form = ServiceForm(request.POST, instance=service)
        if form.is_valid():
            service = form.save()
            
            # Обработка формсета с категориями зрителей
            formset = ViewerCategoryFormset(request.POST, instance=service)
            if formset.is_valid():
                formset.save()
                return redirect('main:create_docs')
    else:
        form = ServiceForm(instance=service)
        formset = ViewerCategoryFormset(instance=service)
    
    return render(request, 'main/create_service.html', {
        'form': form,
        'formset': formset,
        'service': service
    })


@login_required
def document_detail(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    return render(request, 'main/download_docs.html', {'document': document})


@login_required
def download_act(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    # Получаем путь к шаблону акта приемки
    template_path = os.path.join(os.path.dirname(__file__), 'document_templates', 'act_priemki.docx')
    
    # Получаем все услуги
    services = document.services.all()
    # Создаем строку с именами услуг через запятую
    services_text = ", ".join([service.name for service in services])
    
    # Создаем строку с датами через запятую
    dates = ", ".join([service.date for service in services])
    
    # Создаем строку, содержащую название услуги и дату
    service_date = ", ".join([f'{service.name} - {service.date}' for service in services])
    
    # Создаем структуру данных для таблицы с услугами и категориями зрителей
    services_table = []
    total_sum = 0
    total_viewers = 0
    
    for i, service in enumerate(services, 1):
        # Если у услуги есть категории зрителей, создаем строки для каждой категории
        if service.viewer_categories.exists():
            for j, category in enumerate(service.viewer_categories.all(), 1):
                category_total = category.viewers * category.price
                total_sum += category_total
                total_viewers += category.viewers
                
                services_table.append({
                    'num': i if j == 1 else '',  # Номер услуги только для первой категории
                    'service_name': service.name if j == 1 else '',  # Название услуги только для первой категории
                    'service_date': service.date if j == 1 else '',  # Дата услуги только для первой категории
                    'category': True,  # Это категория зрителей
                    'viewers': category.viewers,
                    'price': category.price,
                    'total': category_total,
                    'is_first_category': j == 1,  # Флаг первой категории для форматирования
                    'is_last_category': j == service.viewer_categories.count(),  # Флаг последней категории
                    'category_name': f"{category.viewers} чел. по {category.price} руб."
                })
        else:
            # Если у услуги нет категорий, создаем одну строку
            services_table.append({
                'num': i,
                'service_name': service.name,
                'service_date': service.date,
                'category': False,
                'viewers': 0,
                'price': 0,
                'total': 0
            })
    
    # Создаем список категорий зрителей для всех услуг (для обратной совместимости)
    viewer_categories = []
    for service in services:
        for category in service.viewer_categories.all():
            viewer_categories.append({
                'service_name': service.name,
                'viewers': category.viewers,
                'price': category.price,
                'total': category.viewers * category.price
            })
    
    # Создаем контекст с данными для рендеринга шаблона
    context = {
        'act_and_account_number': document.act_and_account_number,
        'customer': document.customer,
        'doer': document.doer,
        'services': services,
        'services_text': services_text,
        'service_date': service_date,
        'service': services.first() if services.exists() else None,
        'address_and_time': document.address_and_time,
        'price_in_figures': document.price_in_figures,
        'price_in_words': document.price_in_words,
        'date': document.date,
        'dates': dates,
        'viewer_categories': viewer_categories,
        'number_basis_of_the_contract': document.number_basis_of_the_contract,
        'services_table': services_table,  # Новая структура для таблицы
        'services_table_total': total_sum,  # Общая сумма для итоговой строки
        'services_table_total_viewers': total_viewers  # Общее количество зрителей
    }
    
    # Используем DocxTemplate для рендеринга
    doc = DocxTemplate(template_path)
    doc.render(context)
    
    # Сохраняем во временный буфер
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    # Отправляем файл пользователю
    response = HttpResponse(f.read())
    response['Content-Disposition'] = f'attachment; filename=Акт_приемки_{document.act_and_account_number}.docx'
    response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    return response


@login_required
def download_invoice(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    # Получаем путь к шаблону счета
    template_path = os.path.join(os.path.dirname(__file__), 'document_templates', 'schet.docx')
    
    # Получаем все услуги
    services = document.services.all()
    # Создаем строку с именами услуг через запятую
    services_text = ", ".join([service.name for service in services])
    
    # Создаем строку с датами через запятую
    dates = ", ".join([service.date for service in services])
    
    # Создаем строку, содержащую название услуги и дату
    service_date = ", ".join([f'{service.name} - {service.date}' for service in services])
    
    # Создаем структуру данных для таблицы с услугами и категориями зрителей
    services_table = []
    total_sum = 0
    total_viewers = 0
    
    for i, service in enumerate(services, 1):
        # Если у услуги есть категории зрителей, создаем строки для каждой категории
        if service.viewer_categories.exists():
            for j, category in enumerate(service.viewer_categories.all(), 1):
                category_total = category.viewers * category.price
                total_sum += category_total
                total_viewers += category.viewers
                
                services_table.append({
                    'num': i if j == 1 else '',  # Номер услуги только для первой категории
                    'service_name': service.name if j == 1 else '',  # Название услуги только для первой категории
                    'service_date': service.date if j == 1 else '',  # Дата услуги только для первой категории
                    'category': True,  # Это категория зрителей
                    'viewers': category.viewers,
                    'price': category.price,
                    'total': category_total,
                    'is_first_category': j == 1,  # Флаг первой категории для форматирования
                    'is_last_category': j == service.viewer_categories.count(),  # Флаг последней категории
                    'category_name': f"{category.viewers} чел. по {category.price} руб."
                })
        else:
            # Если у услуги нет категорий, создаем одну строку
            services_table.append({
                'num': i,
                'service_name': service.name,
                'service_date': service.date,
                'category': False,
                'viewers': 0,
                'price': 0,
                'total': 0
            })
    
    # Создаем список категорий зрителей для всех услуг (для обратной совместимости)
    viewer_categories = []
    for service in services:
        for category in service.viewer_categories.all():
            viewer_categories.append({
                'service_name': service.name,
                'viewers': category.viewers,
                'price': category.price,
                'total': category.viewers * category.price
            })
    
    # Создаем контекст с данными для рендеринга шаблона
    context = {
        'act_and_account_number': document.act_and_account_number,
        'customer': document.customer,
        'doer': document.doer,
        'services': services,
        'services_text': services_text,
        'service_date': service_date,
        'service': services.first() if services.exists() else None,  # Для совместимости с шаблоном
        'address_and_time': document.address_and_time,
        'price_in_figures': document.price_in_figures,
        'price_in_words': document.price_in_words,
        'date': document.date,
        'time': document.time,
        'dates': dates,
        'viewer_categories': viewer_categories,
        'services_table': services_table,  # Новая структура для таблицы
        'services_table_total': total_sum,  # Общая сумма для итоговой строки
        'services_table_total_viewers': total_viewers,  # Общее количество зрителей
        'number_basis_of_the_contract': document.number_basis_of_the_contract  # Добавляем номер договора
    }
    
    # Используем DocxTemplate для рендеринга
    doc = DocxTemplate(template_path)
    doc.render(context)
    
    # Сохраняем во временный буфер
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    # Отправляем файл пользователю
    response = HttpResponse(f.read())
    response['Content-Disposition'] = f'attachment; filename=Счет_{document.act_and_account_number}.docx'
    response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    return response


@login_required
def download_contract(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    # Получаем путь к шаблону договора
    template_path = os.path.join(os.path.dirname(__file__), 'document_templates', 'dogovor.docx')
    
    # Получаем все услуги
    services = document.services.all()
    # Создаем строку с именами услуг через запятую
    services_text = ", ".join([service.name for service in services])
    
    # Создаем строку с датами через запятую
    dates = ", ".join([service.date for service in services])
    
    # Создаем строку, содержащую название услуги и дату
    service_date = ", ".join([f'{service.name} - {service.date}' for service in services])
    
    # Создаем структуру данных для таблицы с услугами и категориями зрителей
    services_table = []
    total_sum = 0
    total_viewers = 0
    
    for i, service in enumerate(services, 1):
        # Если у услуги есть категории зрителей, создаем строки для каждой категории
        if service.viewer_categories.exists():
            for j, category in enumerate(service.viewer_categories.all(), 1):
                category_total = category.viewers * category.price
                total_sum += category_total
                total_viewers += category.viewers
                
                services_table.append({
                    'num': i if j == 1 else '',  # Номер услуги только для первой категории
                    'service_name': service.name if j == 1 else '',  # Название услуги только для первой категории
                    'service_date': service.date if j == 1 else '',  # Дата услуги только для первой категории
                    'category': True,  # Это категория зрителей
                    'viewers': category.viewers,
                    'price': category.price,
                    'total': category_total,
                    'is_first_category': j == 1,  # Флаг первой категории для форматирования
                    'is_last_category': j == service.viewer_categories.count(),  # Флаг последней категории
                    'category_name': f"{category.viewers} чел. по {category.price} руб."
                })
        else:
            # Если у услуги нет категорий, создаем одну строку
            services_table.append({
                'num': i,
                'service_name': service.name,
                'service_date': service.date,
                'category': False,
                'viewers': 0,
                'price': 0,
                'total': 0
            })
    
    # Создаем список категорий зрителей для всех услуг (для обратной совместимости)
    viewer_categories = []
    for service in services:
        for category in service.viewer_categories.all():
            viewer_categories.append({
                'service_name': service.name,
                'viewers': category.viewers,
                'price': category.price,
                'total': category.viewers * category.price
            })
    
    # Создаем контекст с данными для рендеринга шаблона
    context = {
        'number_basis_of_the_contract': document.number_basis_of_the_contract,
        'customer': document.customer,
        'doer': document.doer,
        'services': services,
        'services_text': services_text,
        'service_date': service_date,
        'service': services.first() if services.exists() else None,  # Для совместимости с шаблоном
        'address_and_time': document.address_and_time,
        'price_in_figures': document.price_in_figures,
        'price_in_words': document.price_in_words,
        'total_viewers': document.viewers,
        'time': document.time,
        'date': document.date,
        'dates': dates,
        'viewer_categories': viewer_categories,
        'services_table': services_table,  # Новая структура для таблицы
        'services_table_total': total_sum,  # Общая сумма для итоговой строки
        'services_table_total_viewers': total_viewers  # Общее количество зрителей
    }
    
    # Используем DocxTemplate для рендеринга
    doc = DocxTemplate(template_path)
    doc.render(context)
    
    # Сохраняем во временный буфер
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    # Отправляем файл пользователю
    response = HttpResponse(f.read())
    response['Content-Disposition'] = f'attachment; filename=Договор_{document.number_basis_of_the_contract}.docx'
    response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    return response


@login_required
def download_all_docs(request, document_id):
    """Функция для скачивания всех документов в одном ZIP-архиве"""
    import zipfile
    from io import BytesIO
    
    document = get_object_or_404(Document, id=document_id)
    
    # Создаем ZIP-архив в памяти
    zip_buffer = BytesIO()
    
    # Создаем безопасное имя папки из short_name клиента для использования внутри архива
    try:
        # Получаем короткое имя организации
        short_name = document.customer.short_name
        # Берём первые 10 символов и убираем небезопасные символы
        folder_name = short_name[:20].strip().replace(' ', '_').replace('"', '').replace("'", "").replace('/', '_')
        # Проверяем, что имя не пустое
        if not folder_name:
            folder_name = f"Документы_{document.act_and_account_number}"
    except (AttributeError, TypeError):
        # Если произошла ошибка (например, short_name не существует)
        folder_name = f"Документы_{document.act_and_account_number}"
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Подготавливаем и добавляем акт приемки
        act_buffer = BytesIO()
        act_template_path = os.path.join(os.path.dirname(__file__), 'document_templates', 'act_priemki.docx')
        
        # Получаем все услуги и готовим данные (как в функции download_act)
        services = document.services.all()
        services_text = ", ".join([service.name for service in services])
        dates = ", ".join([service.date for service in services])
        service_date = ", ".join([f'{service.name} - {service.date}' for service in services])
        
        # Создаем структуру данных для таблицы с услугами и категориями зрителей
        services_table = []
        total_sum = 0
        total_viewers = 0
        
        for i, service in enumerate(services, 1):
            if service.viewer_categories.exists():
                for j, category in enumerate(service.viewer_categories.all(), 1):
                    category_total = category.viewers * category.price
                    total_sum += category_total
                    total_viewers += category.viewers
                    
                    services_table.append({
                        'num': i if j == 1 else '',
                        'service_name': service.name if j == 1 else '',
                        'service_date': service.date if j == 1 else '',
                        'category': True,
                        'viewers': category.viewers,
                        'price': category.price,
                        'total': category_total,
                        'is_first_category': j == 1,
                        'is_last_category': j == service.viewer_categories.count(),
                        'category_name': f"{category.viewers} чел. по {category.price} руб."
                    })
            else:
                services_table.append({
                    'num': i,
                    'service_name': service.name,
                    'service_date': service.date,
                    'category': False,
                    'viewers': 0,
                    'price': 0,
                    'total': 0
                })
        
        # Создаем список категорий зрителей (для обратной совместимости)
        viewer_categories = []
        for service in services:
            for category in service.viewer_categories.all():
                viewer_categories.append({
                    'service_name': service.name,
                    'viewers': category.viewers,
                    'price': category.price,
                    'total': category.viewers * category.price
                })
        
        # Создаем контекст для акта
        act_context = {
            'act_and_account_number': document.act_and_account_number,
            'customer': document.customer,
            'doer': document.doer,
            'services': services,
            'services_text': services_text,
            'service_date': service_date,
            'service': services.first() if services.exists() else None,
            'address_and_time': document.address_and_time,
            'price_in_figures': document.price_in_figures,
            'price_in_words': document.price_in_words,
            'date': document.date,
            'dates': dates,
            'viewer_categories': viewer_categories,
            'number_basis_of_the_contract': document.number_basis_of_the_contract,
            'services_table': services_table,
            'services_table_total': total_sum,
            'services_table_total_viewers': total_viewers
        }
        
        # Рендерим акт и добавляем в архив
        act_doc = DocxTemplate(act_template_path)
        act_doc.render(act_context)
        act_doc.save(act_buffer)
        act_buffer.seek(0)
        zip_file.writestr(f'{folder_name}/Акт_приемки_{document.act_and_account_number}.docx', act_buffer.getvalue())
        
        # Подготавливаем и добавляем счет
        invoice_buffer = BytesIO()
        invoice_template_path = os.path.join(os.path.dirname(__file__), 'document_templates', 'schet.docx')
        
        # Контекст для счета
        invoice_context = {
            'act_and_account_number': document.act_and_account_number,
            'customer': document.customer,
            'doer': document.doer,
            'services': services,
            'services_text': services_text,
            'service_date': service_date,
            'service': services.first() if services.exists() else None,
            'address_and_time': document.address_and_time,
            'price_in_figures': document.price_in_figures,
            'price_in_words': document.price_in_words,
            'date': document.date,
            'time': document.time,
            'dates': dates,
            'viewer_categories': viewer_categories,
            'services_table': services_table,
            'services_table_total': total_sum,
            'services_table_total_viewers': total_viewers,
            'number_basis_of_the_contract': document.number_basis_of_the_contract
        }
        
        # Рендерим счет и добавляем в архив
        invoice_doc = DocxTemplate(invoice_template_path)
        invoice_doc.render(invoice_context)
        invoice_doc.save(invoice_buffer)
        invoice_buffer.seek(0)
        zip_file.writestr(f'{folder_name}/Счет_{document.act_and_account_number}.docx', invoice_buffer.getvalue())
        
        # Подготавливаем и добавляем договор
        contract_buffer = BytesIO()
        contract_template_path = os.path.join(os.path.dirname(__file__), 'document_templates', 'dogovor.docx')
        
        # Контекст для договора
        contract_context = {
            'number_basis_of_the_contract': document.number_basis_of_the_contract,
            'customer': document.customer,
            'doer': document.doer,
            'services': services,
            'services_text': services_text,
            'service_date': service_date,
            'service': services.first() if services.exists() else None,
            'address_and_time': document.address_and_time,
            'price_in_figures': document.price_in_figures,
            'price_in_words': document.price_in_words,
            'total_viewers': document.viewers,
            'time': document.time,
            'date': document.date,
            'dates': dates,
            'viewer_categories': viewer_categories,
            'services_table': services_table,
            'services_table_total': total_sum,
            'services_table_total_viewers': total_viewers
        }
        
        # Рендерим договор и добавляем в архив
        contract_doc = DocxTemplate(contract_template_path)
        contract_doc.render(contract_context)
        contract_doc.save(contract_buffer)
        contract_buffer.seek(0)
        zip_file.writestr(f'{folder_name}/Договор_{document.number_basis_of_the_contract}.docx', contract_buffer.getvalue())
    
    # Готовим ZIP-архив для отправки
    zip_buffer.seek(0)
    
    # Создаем безопасное имя файла для zip-архива (аналогично имени папки внутри)
    safe_name = folder_name
    
    # Отправляем архив пользователю
    response = HttpResponse(zip_buffer.getvalue())
    response['Content-Type'] = 'application/zip'
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.zip"'
    
    return response


@login_required
def service_ajax(request):
    """Функция для обработки AJAX-запросов для создания/редактирования/удаления услуг"""
    if request.method == 'POST':
        try:
            # Чтение и декодирование JSON-данных из запроса
            data = json.loads(request.body)
            action = data.get('action')
            
            if action == 'create':
                # Создание новой услуги
                service_data = data.get('service', {})
                form = ServiceForm(service_data)
                if form.is_valid():
                    service = form.save()
                    
                    # Обработка категорий зрителей
                    categories_data = data.get('categories', [])
                    for category_data in categories_data:
                        ViewerCategory.objects.create(
                            service=service,
                            viewers=category_data.get('viewers', 0),
                            price=category_data.get('price', 0)
                        )
                    
                    return JsonResponse({
                        'success': True,
                        'service': {
                            'id': service.id,
                            'name': service.name,
                            'date': service.date,
                            'total_viewers': service.total_viewers,
                            'total_price': service.total_price,
                            'categories': [
                                {
                                    'id': category.id,
                                    'viewers': category.viewers,
                                    'price': category.price
                                } for category in service.viewer_categories.all()
                            ]
                        }
                    })
                else:
                    return JsonResponse({'success': False, 'errors': form.errors})
            
            elif action == 'edit':
                # Редактирование существующей услуги
                service_id = data.get('service_id')
                service = get_object_or_404(Service, id=service_id)
                
                service_data = data.get('service', {})
                form = ServiceForm(service_data, instance=service)
                if form.is_valid():
                    service = form.save()
                    
                    # Удаляем существующие категории
                    service.viewer_categories.all().delete()
                    
                    # Добавляем новые категории
                    categories_data = data.get('categories', [])
                    for category_data in categories_data:
                        ViewerCategory.objects.create(
                            service=service,
                            viewers=category_data.get('viewers', 0),
                            price=category_data.get('price', 0)
                        )
                    
                    return JsonResponse({
                        'success': True,
                        'service': {
                            'id': service.id,
                            'name': service.name,
                            'date': service.date,
                            'total_viewers': service.total_viewers,
                            'total_price': service.total_price,
                            'categories': [
                                {
                                    'id': category.id,
                                    'viewers': category.viewers,
                                    'price': category.price
                                } for category in service.viewer_categories.all()
                            ]
                        }
                    })
                else:
                    return JsonResponse({'success': False, 'errors': form.errors})
            
            elif action == 'delete':
                # Удаление услуги
                service_id = data.get('service_id')
                service = get_object_or_404(Service, id=service_id)
                service.delete()
                return JsonResponse({'success': True})
            
            else:
                return JsonResponse({'success': False, 'error': 'Неизвестное действие'})
                
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Неверный формат JSON данных'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Метод не поддерживается'})


@login_required
def extract_organisation_data(request):
    """Функция для обработки загруженного документа и извлечения данных организации с помощью ИИ"""
    if request.method == 'POST':
        try:
            # Получаем файл из запроса
            uploaded_file = request.FILES.get('document')
            
            if not uploaded_file:
                return JsonResponse({'success': False, 'error': 'Файл не загружен'})
            
            # Проверяем, что файл имеет расширение .docx
            if not uploaded_file.name.endswith('.docx'):
                return JsonResponse({'success': False, 'error': 'Загруженный файл должен быть в формате DOCX'})
            
            # Получаем содержимое документа
            document_text = extract_text_from_docx(uploaded_file)
            
            # Если текст успешно извлечен, отправляем его в ИИ для анализа
            if document_text:
                organisation_data = extract_organisation_data_with_ai(document_text)
                
                return JsonResponse({
                    'success': True,
                    'data': organisation_data
                })
            else:
                return JsonResponse({'success': False, 'error': 'Не удалось извлечь текст из документа'})
                
        except Exception as e:
            print(f"Error processing document: {e}")
            return JsonResponse({'success': False, 'error': str(e)})
    else:
        # Если метод не POST, возвращаем ошибку
        return JsonResponse({'success': False, 'error': 'Метод не поддерживается'})


@login_required
def get_organisations_data(request):
    """Функция для получения данных организаций (адресов) для автозаполнения"""
    # Собираем данные всех организаций в словарь
    organisations_data = {}
    
    # Получаем все организации
    organisations = Organisation.objects.all()
    
    # Заполняем словарь данными
    for org in organisations:
        organisations_data[str(org.id)] = {
            'address': org.address,
            'short_name': org.short_name
        }
    
    # Возвращаем данные в формате JSON
    return JsonResponse(organisations_data)


def extract_text_from_docx(docx_file):
    """Функция для извлечения текста из .docx файла"""
    try:
        # Создаем временный файл для сохранения загруженного файла
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            for chunk in docx_file.chunks():
                temp_file.write(chunk)
            temp_file_path = temp_file.name
        
        # Открываем документ с помощью python-docx
        doc = DocxDocument(temp_file_path)
        
        # Извлекаем весь текст из документа
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        # Извлекаем текст из таблиц, если они есть
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text)
        
        # Удаляем временный файл
        os.unlink(temp_file_path)
        
        # Возвращаем весь текст, объединенный символами новой строки
        return '\n'.join(full_text)
        
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None


def extract_organisation_data_with_ai(document_text):
    """Функция для извлечения информации об организации из текста с использованием ИИ"""
    try:
        # Инициализируем клиент OpenAI (используем VseGPT)
        client = OpenAI(
            api_key=os.environ.get("VSEGPT_API_KEY", ""),
            base_url="https://api.vsegpt.ru/v1",
        )
        
        # Подготавливаем промпт для ИИ с инструкциями
        prompt = f"""
        Ты - помощник по извлечению реквизитов организации из документов.
        
        Ниже приведен текст из договора или иного документа. Мне нужно извлечь из него данные об организации-контрагенте (обычно "Заказчик").
        
        Вот список полей, которые нужно заполнить:
        - name: полное название организации с представителем в формате "Полное название организации, в лице должность ФИО". Например: "Муниципальное бюджетное общеобразовательное учреждение «Гимназия №54», в лице директора Исаева Рамиля Робертовича"
        - short_name: краткое название организации без представителя. Например: "МБОУ «Гимназия №54»"
        - fio: ФИО руководителя. Например: "Исаев Рамиль Робертович"
        - function: должность руководителя (только "Директор" или "Индивидуальный предприниматель")
        - inn: ИНН организации (только цифры)
        - kpp: КПП организации (только цифры)
        - bik: БИК банка (только цифры)
        - bank_name: название банка
        - correspondent_bank_account: корреспондентский счет банка (только цифры)
        - bank_account: расчетный счет организации (только цифры)
        - address: юридический адрес организации
        - requisites: полные реквизиты организации, отформатированные в виде текста с переносами строк
        
        В поле name ОБЯЗАТЕЛЬНО включить представителя организации в формате: "Название, в лице должность ФИО".
        
        Для поля requisites сформируй текст в следующем формате с переносами строк:
        ```
        [Краткое название организации]
        
        [Адрес организации (в одну или две строки)]
        ИНН [ИНН], КПП [КПП],
        р/с [Расчетный счет],
        к/с [Корреспондентский счет],
        [Дополнительные счета, если есть],
        банк: [Название банка],
        БИК: [БИК].
        
        [Дополнительная контактная информация, если есть]
        ```
        
        Пример правильного форматирования поля requisites:
        ```
        Муниципальное бюджетное общеобразовательное учреждение «Гимназия №54»
        
        423832, РТ, г. Набережные Челны, 
        Бульвар Касима, д. 6   
        ИНН 1650081400, КПП 165001001,
        р/с 03234643927300001100,
        к/с 40102810445370000079,
        л/с ЛБГ30800667 Гим54 (бюджетный),
        л/с ЛБВ30800668 Гим54 (внебюджетный),      
        л/с ЛБО30800721 Гим54 (иные цели),
        л/с ЛР308000149-Гим54  (обеспечение по р/с 03234643927300001101),
        банк: ОТДЕЛЕНИЕ – НБ РЕСПУБЛИКА ТАТАРСТАН БАНК РОССИИ/УФК по Республике Татарстан г.Казань,
        БИК: 019205400.
        
        Тел. (8552) 34-03-37, 34-56-80
        Е-Mail: gim54_chelny@mail.ru
        ```
        
        Отвечай ТОЛЬКО в формате JSON, чтобы я мог сразу использовать ответ в своей программе. Если какое-то поле найти не удалось, оставь его пустым.
        
        Вот текст документа:
        {document_text}
        """
        
        messages = []
        messages.append({"role": "user", "content": prompt})
        
        # Отправляем запрос к API
        response = client.chat.completions.create(
            model="anthropic/claude-3-haiku",
            messages=messages,
            temperature=0.2,
            max_tokens=3000,
            extra_headers={"X-Title": "Organisation Extractor"},
        )
        
        # Получаем ответ
        ai_response = response.choices[0].message.content
        
        # Очищаем ответ от markdown-обертки, если она есть
        ai_response = re.sub(r'```json', '', ai_response)
        ai_response = re.sub(r'```', '', ai_response)
        
        # Парсим JSON из ответа ИИ
        organisation_data = json.loads(ai_response.strip())
        
        # Добавляем проверку на наличие всех необходимых полей
        required_fields = [
            'name', 'short_name', 'fio', 'function', 'inn', 'kpp', 'bik', 
            'bank_name', 'correspondent_bank_account', 'bank_account', 'address', 'requisites'
        ]
        
        for field in required_fields:
            if field not in organisation_data:
                organisation_data[field] = ""
        
        return organisation_data
        
    except Exception as e:
        print(f"Error in AI extraction: {e}")
        # Возвращаем пустой словарь в случае ошибки
        return {
            'name': "", 'short_name': "", 'fio': "", 'function': "Директор",
            'inn': "", 'kpp': "", 'bik': "", 'bank_name': "",
            'correspondent_bank_account': "", 'bank_account': "", 'address': "", 'requisites': ""
        }